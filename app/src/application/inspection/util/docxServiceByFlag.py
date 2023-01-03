import math
from datetime import datetime

from application.core import config
from application.inspection.models.inspectionPdf import (InspectionPdfDTO,
                                                         Inspector)
from application.inspection.models.pdfPhoto import PdfPhoto
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas
from textwrap import wrap

from docx import Document
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

MPMG_logo = ImageReader('../assets/mpmg_pdf_logo.png')

PDF_SIDE_SPACING = 24.45
PDF_TOP_SPACING = 8.10

PDF_SIDE_INITIAL_VALUE = 25
PDF_TOP_INITIAL_VALUE = 100
PDF_BOTTOM_LIMIT = 160

def __mm_to_p(mm):
  return mm / 0.352777


def __draw_template(pdf: Canvas):
  timestamp = datetime.now().strftime("%d/%m/%Y às %H:%M:%S")
  pdf.setFont("Helvetica-Bold", 14)
  pdf.drawString(__mm_to_p(60), __mm_to_p(250), "Relatório Automático de Vistoria")
  pdf.setFont("Helvetica", 12)
  pdf.drawString(__mm_to_p(30), __mm_to_p(235), f"Software: {config.settings.system_name}")
  pdf.drawString(__mm_to_p(30), __mm_to_p(225), f"Versão: {config.settings.system_version}")
  pdf.drawString(__mm_to_p(30), __mm_to_p(215), f"Relatório gerado em: {timestamp}")


def __draw_header(pdf: Canvas):
  pdf.drawImage(MPMG_logo,__mm_to_p(PDF_SIDE_INITIAL_VALUE), __mm_to_p(265) , width=150, height=60)
  pdf.setFont("Helvetica-Bold", 12)
  pdf.drawString(__mm_to_p(110), __mm_to_p(280) , config.settings.report_institution)
  pdf.drawString(__mm_to_p(110), __mm_to_p(275), config.settings.report_department)
  pdf.drawString(__mm_to_p(110), __mm_to_p(270), config.settings.report_section)
  pdf.setFont("Helvetica", 12)


def __draw_footer(pdf: Canvas):
  page_num = pdf.getPageNumber()
  pdf.setFont("Helvetica-Bold", 8)
  pdf.drawString(__mm_to_p(160), __mm_to_p(21), f'Página {page_num}')
  pdf.line(__mm_to_p(20), __mm_to_p(20), __mm_to_p(180), __mm_to_p(20))
  pdf.setFont("Helvetica", 9)
  pdf.drawString(__mm_to_p(45), __mm_to_p(15), config.settings.report_address)
  pdf.drawString(__mm_to_p(70), __mm_to_p(10), config.settings.report_contact)
  pdf.setFont("Helvetica-Bold", 9)
  pdf.drawString(__mm_to_p(148), __mm_to_p(10), config.settings.report_website)

  link_rect = (__mm_to_p(148), __mm_to_p(10), __mm_to_p(148), __mm_to_p(10))
  pdf.linkURL(f'http://{config.settings.report_website}', rect=link_rect, relative=1)
  pdf.setFont("Helvetica", 12)


def __draw_content(pdf: Canvas, position: tuple, content: PdfPhoto, index: int):
  width = 400
  height = 200
  x, y = position
  image_path = getattr(content, 'image_path')
  image = ImageReader(image_path)
  end_point_x = x
  end_point_y = y - 40
  latitude = getattr(content, 'latitude')
  longitude = getattr(content, 'longitude')
  description = getattr(content, 'description') if getattr(content, 'description') is not None else "Não há descrição"

  pdf.setFont("Helvetica", 9)
  pdf.rect(x, y, width, height)
  pdf.drawImage(image, x, y, width, height)
  pdf.linkURL(f'http://0.0.0.0:8000/images/{image_path}', rect=(x, y, width+x, height+y), relative=1)
  pdf.rect(end_point_x, end_point_y, width, height/5)
  pdf.drawString(x + 4, end_point_y + 6, f"Coordenadas geográficas: {__convert_coordinates_to_DMS(latitude, longitude)}")
  pdf.drawString(x + 4, end_point_y + 28, f"Foto {index}: {description}")

  pdf.drawString(x + 250, end_point_y + 6, '(ver no mapa)')
  link_rect = (x + 250, end_point_y + 4, (x+250) + 55, (end_point_y + 6) + 10)
  pdf.linkURL(f'https://maps.google.com/?q={latitude},{longitude}&z=15', rect=link_rect, relative=1)

  pdf.setFont("Helvetica", 12)

  return end_point_x, end_point_y - 215


def __draw_comments(pdf: Canvas, observations: str, inspector: Inspector):
  pdf.showPage()
  __draw_header(pdf)
  __draw_footer(pdf)
  pdf.setFont("Helvetica", 12)
  pdf.drawString(__mm_to_p(30), __mm_to_p(235), 'Observações:')

  # Handle multi-linetext
  textobject = pdf.beginText()
  textobject.setTextOrigin(__mm_to_p(30), __mm_to_p(220))
  observations_with_line_breaks = "\n".join(wrap(observations, 80)) # 80 is line width
  textobject.textLines(observations_with_line_breaks)
  pdf.drawText(textobject)

  pdf.line(__mm_to_p(50), __mm_to_p(140), __mm_to_p(150), __mm_to_p(140))
  pdf.drawString(__mm_to_p(70), __mm_to_p(135), f"{getattr(inspector, 'name')} - {getattr(inspector, 'email')}")
  pdf.drawString(__mm_to_p(70), __mm_to_p(130), f"{getattr(inspector, 'role')}")


def __getDMS(value: str):
  value =  abs(float(value))
  degree = math.floor(value)
  minutes = math.floor((value - degree) * 60)
  seconds = round((value - degree - minutes / 60) * 3600 * 1000) / 1000

  return f"{degree}º{minutes}'{seconds}\""


def __convert_coordinates_to_DMS(lat: str, lng: str):
  finalLat = __getDMS(lat)
  finalLat += 'N' if float(lat) >= 0 else 'S'
  finalLng = __getDMS(lng)
  finalLng += 'L' if float(lat) >= 0 else 'O'

  return finalLat + ", " + finalLng


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def __draw_content(document: Document, content: PdfPhoto, index: int):
  image_path = getattr(content, 'image_path')
  latitude = getattr(content, 'latitude')
  longitude = getattr(content, 'longitude')
  description = getattr(content, 'description') if getattr(content, 'description') is not None else "Não há descrição"

  document.add_picture(image_path, width=Inches(5.75))
  pDescription = document.add_paragraph(f'Foto {index+1}: {description} ')
  add_hyperlink(pDescription, f'http://0.0.0.0:8000/images/{image_path}', "(download)", '3366CC', True)
  pCoordinates = document.add_paragraph(f'Coordenadas geográficas: {__convert_coordinates_to_DMS(latitude, longitude)} ')
  add_hyperlink(pCoordinates, f'https://maps.google.com/?q={latitude},{longitude}&z=15', "(ver no mapa)", '3366CC', True)


def generate_docx_by_flag(data: InspectionPdfDTO):
  timestamp = datetime.now().strftime("%d/%m/%Y às %H:%M:%S")
  document = Document()
  style = document.styles['Normal']
  font = style.font
  font.name = 'Helvetica'
  font.size = Pt(12)

  # header = document.sections[0].header
  # paragraph = header.paragraphs[0]

  # logo_run = paragraph.add_run()
  # logo_run.add_picture('../assets/mpmg_pdf_logo.png', width=Inches(1))

  # text_run = paragraph.add_run()
  # text_run.text = '\t\t' + f'{config.settings.report_institution}\n\t\t{config.settings.report_department}\n\t\t{config.settings.report_section}' # For center align of text

  header = document.sections[0].header
  htable=header.add_table(1, 2, Inches(6))
  htable.style = 'Table Grid'
  htab_cells=htable.rows[0].cells
  ht0=htab_cells[0].add_paragraph()
  kh=ht0.add_run()
  kh.add_picture('../assets/mpmg_pdf_logo.png', width=Inches(1.5))
  ht0=htab_cells[0].add_paragraph()
  ht1=htab_cells[1].add_paragraph(f'{config.settings.report_institution}\n{config.settings.report_department}\n{config.settings.report_section}')
  ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

  footer = document.sections[0].footer
  # Calling the paragraph already present in
  # the footer section
  footer_para = footer.paragraphs[0]
  footerRun = footer_para.add_run(f"______________________________________________________________________________________\n{config.settings.report_address}\n{config.settings.report_contact} {config.settings.report_website}")
  footerRun.font.size = Pt(9)
  footerRun.alignment = WD_ALIGN_PARAGRAPH.RIGHT

  # Adding text in the footer
  # footer_para.text = f"________________________________________________________________\n{config.settings.report_address}\n{config.settings.report_contact} {config.settings.report_website}"
  # footer_para.style = "Heading 9"

  pHeader = document.add_paragraph()
  pHeader.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  pHeader.add_run(f'Relatório automático de vistoria').bold = True
  document.add_paragraph(f"Software: {config.settings.system_name}")
  document.add_paragraph(f"Versão: {config.settings.system_version}")
  document.add_paragraph(f"Relatório gerado em: {timestamp}")
  pInquiry = document.add_paragraph()
  pInquiry.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  pInquiry.add_run(f"Vistoria n° {getattr(data, 'inquiry_number')}").bold = True
  document.add_paragraph(f"Local: {getattr(data, 'local')}")
  document.add_paragraph(f"Data da Vistoria: {getattr(data, 'inspection_date')}")

  # p = document.add_paragraph('A plain paragraph having some ')
  # p.add_run('bold').bold = True
  # p.add_run(' and some ')
  # p.add_run('italic.').italic = True

  # document.add_heading('Heading, level 1', level=1)
  # document.add_paragraph('Intense quote', style='Intense Quote')

  # document.add_paragraph(
  #     'first item in unordered list', style='List Bullet'
  # )
  # document.add_paragraph(
  #     'first item in ordered list', style='List Number'
  # )

  for index, content in enumerate(getattr(data, 'content')):
    __draw_content(document, content, index)

  document.add_paragraph()
  pComments = document.add_paragraph()
  pComments.add_run(f"Comentários gerais:").bold = True
  document.add_paragraph(getattr(data, 'observations'))

  inspector = getattr(data, 'inspector')
  document.add_paragraph()
  pSignature = document.add_paragraph("______________________________________")
  pSignature.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  pInspectorName = document.add_paragraph(f"{getattr(inspector, 'name')} - {getattr(inspector, 'email')}")
  pInspectorName.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  pInspectorRole = document.add_paragraph(getattr(inspector, 'role'))
  pInspectorRole.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
  
  # document.add_page_break()

  document.save(f"../reports/relatorio-vistoria-{getattr(data, 'inquiry_number')}.docx")
  return f"../reports/relatorio-vistoria-{getattr(data, 'inquiry_number')}.docx"
  # pdf = Canvas(f"../reports/relatorio-vistoria-{getattr(data, 'inquiry_number')}.pdf", pagesize=A4)
  # pdf.setAuthor("Equipe F05")
  # pdf.setTitle(f"Relatório Automático - Vistoria {getattr(data, 'inquiry_number')}")
  # __draw_header(pdf)
  # __draw_footer(pdf)
  # __draw_template(pdf)

  # pdf.setFont("Helvetica-Bold", 14)
  # pdf.drawString(__mm_to_p(90), __mm_to_p(200), f"Vistoria n° {getattr(data, 'inquiry_number')}")
  # pdf.setFont("Helvetica", 12)

  # pdf.drawString(__mm_to_p(30), __mm_to_p(185), f"Local: {getattr(data, 'local')}")
  # pdf.drawString(__mm_to_p(30), __mm_to_p(175), f"Data da Vistoria: {getattr(data, 'inspection_date')}")
  # entry_point = (__mm_to_p(30), __mm_to_p(90))
  # end_point = (0, 0)

  # for index, content in enumerate(getattr(data, 'content')):
  #   if end_point[1] <= PDF_BOTTOM_LIMIT and end_point[1] != 0: # Reset axis on new page
  #     pdf.showPage()
  #     __draw_header(pdf)
  #     __draw_footer(pdf)
  #     entry_point = (__mm_to_p(30), __mm_to_p(170))
  #     end_point = __draw_content(pdf, entry_point, content, index+1)
  #     entry_point = end_point
  #   else:
  #     end_point = __draw_content(pdf, entry_point, content, index+1)
  #     entry_point = end_point

  # __draw_comments(pdf, getattr(data, 'observations'), getattr(data, 'inspector'))
  # pdf.save()

  # return f"../reports/relatorio-vistoria-{getattr(data, 'inquiry_number')}.pdf"
